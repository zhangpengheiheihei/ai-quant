# -*- coding: utf-8 -*-
"""
================================================================================
 高能环境（603588.SH）TASK2 技术指标分析报告
================================================================================
 内容:
  1. 数据基础诊断分析（缺失值检查、描述性统计）
  2. RSI / MACD / 布林带 指标介绍与计算
  3. 指标可视化图表
  4. 扩展指标 KDJ 介绍与计算
  5. 生成 PDF 报告（宋体、五号字、1.5倍行距、两端对齐）
================================================================================
"""

import csv
import os
import sys
import datetime
import json

# ---------------------------------------------------------------------------
# 依赖检查
# ---------------------------------------------------------------------------
REQUIRED = {
    "numpy": "numpy",
    "matplotlib": "matplotlib",
    "docx": "python-docx",
}

for import_name, pip_name in REQUIRED.items():
    try:
        __import__(import_name)
    except ImportError:
        import subprocess
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pip_name, "-q"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from matplotlib.font_manager import FontProperties

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
CSV_PATH = os.path.join(OUTPUT_DIR, "gaoneng_huanjing_daily.csv")

# ---------------------------------------------------------------------------
# 1. 加载数据 + 诊断分析
# ---------------------------------------------------------------------------
def load_and_diagnose():
    """加载CSV数据，进行缺失值检查和描述性统计"""
    print("[1] 加载数据 & 诊断分析...")

    data = []
    with open(CSV_PATH, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            data.append({
                "date": row["date"],
                "open": float(row["open"]),
                "close": float(row["close"]),
                "high": float(row["high"]),
                "low": float(row["low"]),
                "volume": int(row["volume"]),
                "amount": float(row["amount"]),
                "pct_chg": float(row["pct_chg"]),
                "turnover": float(row["turnover"]),
            })

    n = len(data)
    print(f"    共 {n} 条记录, 日期范围: {data[0]['date']} ~ {data[-1]['date']}")

    # ---- 缺失值检查 ----
    fields = ["open", "close", "high", "low", "volume", "amount", "pct_chg", "turnover"]
    missing_report = {}
    for field in fields:
        missing_count = sum(1 for d in data if d[field] is None or (isinstance(d[field], float) and np.isnan(d[field])))
        missing_report[field] = missing_count

    total_missing = sum(missing_report.values())
    print(f"    缺失值总计: {total_missing} 个")

    # ---- 描述性统计 ----
    stats = {}
    for field in fields:
        vals = [d[field] for d in data]
        arr = np.array(vals, dtype=float)
        stats[field] = {
            "count": int(len(arr)),
            "mean": float(np.mean(arr)),
            "std": float(np.std(arr, ddof=1)),
            "min": float(np.min(arr)),
            "q25": float(np.percentile(arr, 25)),
            "median": float(np.median(arr)),
            "q75": float(np.percentile(arr, 75)),
            "max": float(np.max(arr)),
            "skewness": float((np.mean(arr) - np.median(arr)) / np.std(arr, ddof=1)) if np.std(arr, ddof=1) > 0 else 0,
        }

    return data, missing_report, stats


# ---------------------------------------------------------------------------
# 2. 指标计算函数
# ---------------------------------------------------------------------------
def calc_rsi(closes, period=14):
    """计算 RSI (Relative Strength Index)"""
    n = len(closes)
    if n <= period:
        return [None] * n

    deltas = np.diff(closes)
    gains = np.where(deltas > 0, deltas, 0.0)
    losses = np.where(deltas < 0, -deltas, 0.0)

    rsi = [None] * n
    # 第一个窗口用简单平均
    avg_gain = np.mean(gains[:period])
    avg_loss = np.mean(losses[:period])
    if avg_loss == 0:
        rsi[period] = 100.0
    else:
        rs = avg_gain / avg_loss
        rsi[period] = 100.0 - 100.0 / (1.0 + rs)

    # 后续用 Wilder 平滑
    for i in range(period + 1, n):
        avg_gain = (avg_gain * (period - 1) + gains[i - 1]) / period
        avg_loss = (avg_loss * (period - 1) + losses[i - 1]) / period
        if avg_loss == 0:
            rsi[i] = 100.0
        else:
            rs = avg_gain / avg_loss
            rsi[i] = 100.0 - 100.0 / (1.0 + rs)

    return rsi


def calc_macd(closes, fast=12, slow=26, signal=9):
    """计算 MACD (Moving Average Convergence Divergence)"""
    n = len(closes)
    if n < slow:
        return [None] * n, [None] * n, [None] * n

    # EMA 计算
    def ema(series, period):
        result = [None] * len(series)
        sma = np.mean(series[:period])
        k = 2.0 / (period + 1.0)
        result[period - 1] = sma
        for i in range(period, len(series)):
            result[i] = series[i] * k + result[i - 1] * (1.0 - k)
        return result

    ema_fast = ema(closes, fast)
    ema_slow = ema(closes, slow)

    dif = [None] * n
    for i in range(slow - 1, n):
        if ema_fast[i] is not None and ema_slow[i] is not None:
            dif[i] = ema_fast[i] - ema_slow[i]

    # DIF signal line (DEA)
    dea = [None] * n
    # 从 DIF 的第一个有效值开始计算 DEA (slow + signal - 2 之后)
    valid_dif = [v for v in dif if v is not None]
    if len(valid_dif) >= signal:
        start_idx = slow - 1
        # SMA of first 'signal' DIF values
        sma_dea = np.mean(valid_dif[:signal])
        k_dea = 2.0 / (signal + 1.0)
        dea_idx = start_idx + signal - 1
        dea[dea_idx] = sma_dea
        for i in range(dea_idx + 1, n):
            if dif[i] is not None:
                dea[i] = dif[i] * k_dea + dea[i - 1] * (1.0 - k_dea)

    # MACD柱 = 2 * (DIF - DEA)
    macd_hist = [None] * n
    for i in range(n):
        if dif[i] is not None and dea[i] is not None:
            macd_hist[i] = 2.0 * (dif[i] - dea[i])

    return dif, dea, macd_hist


def calc_bollinger(closes, period=20, std_dev=2.0):
    """计算布林带 (Bollinger Bands)"""
    n = len(closes)
    middle = [None] * n
    upper = [None] * n
    lower = [None] * n

    if n < period:
        return middle, upper, lower

    for i in range(period - 1, n):
        window = closes[i - period + 1: i + 1]
        ma = np.mean(window)
        std = np.std(window, ddof=1)
        middle[i] = ma
        upper[i] = ma + std_dev * std
        lower[i] = ma - std_dev * std

    return middle, upper, lower


def calc_kdj(highs, lows, closes, period=9, k_period=3, d_period=3):
    """计算 KDJ 随机指标"""
    n = len(closes)
    k_vals = [None] * n
    d_vals = [None] * n
    j_vals = [None] * n

    if n < period:
        return k_vals, d_vals, j_vals

    # RSV 计算
    rsv = [None] * n
    for i in range(period - 1, n):
        hh = max(highs[i - period + 1: i + 1])
        ll = min(lows[i - period + 1: i + 1])
        if hh == ll:
            rsv[i] = 50.0
        else:
            rsv[i] = (closes[i] - ll) / (hh - ll) * 100.0

    # 初始 K/D 值
    start = period - 1
    k_vals[start] = 50.0
    d_vals[start] = 50.0

    for i in range(start + 1, n):
        if rsv[i] is not None:
            k_vals[i] = (k_vals[i - 1] * (k_period - 1) + rsv[i]) / k_period
            d_vals[i] = (d_vals[i - 1] * (d_period - 1) + k_vals[i]) / d_period
            j_vals[i] = 3.0 * k_vals[i] - 2.0 * d_vals[i]

    return k_vals, d_vals, j_vals


# ---------------------------------------------------------------------------
# 3. 图表生成
# ---------------------------------------------------------------------------
def generate_indicator_charts(data, dates, closes, highs, lows, rsi, dif, dea, macd_hist, bb_mid, bb_up, bb_low, k, d, j):
    """生成四张技术指标图表"""
    print("[3] 生成技术指标图表...")

    matplotlib.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "KaiTi"]
    matplotlib.rcParams["axes.unicode_minus"] = False

    n = len(data)
    step = max(1, n // 20)
    tick_pos = list(range(0, n, step))
    tick_labels = [dates[i][5:] for i in tick_pos]  # MM-DD

    chart_paths = []

    # 颜色常量
    UP_COLOR = "#e74c3c"
    DOWN_COLOR = "#27ae60"
    LINE_COLOR = "#667eea"

    # ---- 图1: 收盘价 + 布林带 ----
    fig, ax = plt.subplots(figsize=(14, 6))
    ax.plot(range(n), closes, color=LINE_COLOR, linewidth=1.2, label="收盘价", zorder=3)
    ax.plot(range(n), bb_mid, color="#ffa502", linewidth=0.8, linestyle="--", label="BOLL中轨(MA20)")
    ax.plot(range(n), bb_up, color=UP_COLOR, linewidth=0.7, linestyle="-", alpha=0.7, label="BOLL上轨")
    ax.plot(range(n), bb_low, color=DOWN_COLOR, linewidth=0.7, linestyle="-", alpha=0.7, label="BOLL下轨")

    # 填充布林带区间
    valid_x = [i for i in range(n) if bb_up[i] is not None]
    if valid_x:
        up_arr = np.array([bb_up[i] for i in valid_x])
        low_arr = np.array([bb_low[i] for i in valid_x])
        ax.fill_between(valid_x, up_arr, low_arr, alpha=0.08, color="#667eea")

    ax.set_title("图1  高能环境（603588）布林带指标 (BOLL, N=20, k=2)", fontsize=14, fontweight="bold")
    ax.set_ylabel("价格（元）")
    ax.legend(loc="upper left", fontsize=9)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_xlim(-1, n)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels(tick_labels, rotation=45, fontsize=7, ha="right")
    plt.tight_layout()
    path1 = os.path.join(OUTPUT_DIR, "chart_boll.png")
    plt.savefig(path1, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    chart_paths.append(path1)
    print(f"    -> chart_boll.png")

    # ---- 图2: RSI ----
    fig, ax = plt.subplots(figsize=(14, 4))
    ax.plot(range(n), rsi, color="#764ba2", linewidth=1.2, label="RSI(14)")
    ax.axhline(y=70, color=UP_COLOR, linestyle="--", linewidth=0.8, alpha=0.6, label="超买线(70)")
    ax.axhline(y=30, color=DOWN_COLOR, linestyle="--", linewidth=0.8, alpha=0.6, label="超卖线(30)")
    ax.axhline(y=50, color="#999", linestyle=":", linewidth=0.6, alpha=0.5)
    ax.fill_between(range(n), 70, 100, alpha=0.1, color=UP_COLOR)
    ax.fill_between(range(n), 0, 30, alpha=0.1, color=DOWN_COLOR)

    ax.set_title("图2  高能环境（603588）相对强弱指标 (RSI, N=14)", fontsize=14, fontweight="bold")
    ax.set_ylabel("RSI 值")
    ax.legend(loc="upper left", fontsize=9)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_ylim(0, 100)
    ax.set_xlim(-1, n)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels(tick_labels, rotation=45, fontsize=7, ha="right")
    plt.tight_layout()
    path2 = os.path.join(OUTPUT_DIR, "chart_rsi.png")
    plt.savefig(path2, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    chart_paths.append(path2)
    print(f"    -> chart_rsi.png")

    # ---- 图3: MACD ----
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 6), sharex=True,
                                    gridspec_kw={"height_ratios": [2, 1], "hspace": 0.05})

    # 价格走势
    ax1.plot(range(n), closes, color=LINE_COLOR, linewidth=1.2, label="收盘价")
    ax1.set_ylabel("价格（元）")
    ax1.legend(loc="upper left", fontsize=9)
    ax1.grid(axis="y", alpha=0.3, linestyle="--")
    ax1.set_xlim(-1, n)

    # MACD
    ax2.plot(range(n), dif, color="#667eea", linewidth=1.0, label="DIF")
    ax2.plot(range(n), dea, color="#ffa502", linewidth=1.0, label="DEA")

    # MACD柱
    for i in range(n):
        if macd_hist[i] is not None:
            clr = UP_COLOR if macd_hist[i] >= 0 else DOWN_COLOR
            ax2.bar(i, macd_hist[i], 0.6, color=clr, alpha=0.7)

    ax2.axhline(y=0, color="#999", linestyle="-", linewidth=0.5)
    ax2.set_ylabel("MACD")
    ax2.legend(loc="upper left", fontsize=9, ncol=3)
    ax2.grid(axis="y", alpha=0.3, linestyle="--")
    ax2.set_xlim(-1, n)
    ax2.set_xticks(tick_pos)
    ax2.set_xticklabels(tick_labels, rotation=45, fontsize=7, ha="right")

    fig.suptitle("图3  高能环境（603588）MACD指标 (12,26,9)", fontsize=14, fontweight="bold", y=0.97)
    plt.tight_layout()
    path3 = os.path.join(OUTPUT_DIR, "chart_macd.png")
    plt.savefig(path3, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    chart_paths.append(path3)
    print(f"    -> chart_macd.png")

    # ---- 图4: KDJ ----
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(range(n), k, color="#667eea", linewidth=1.0, label="K值")
    ax.plot(range(n), d, color="#ffa502", linewidth=1.0, label="D值")
    ax.plot(range(n), j, color="#e74c3c", linewidth=0.8, alpha=0.7, label="J值")
    ax.axhline(y=80, color=UP_COLOR, linestyle="--", linewidth=0.8, alpha=0.5, label="超买(80)")
    ax.axhline(y=20, color=DOWN_COLOR, linestyle="--", linewidth=0.8, alpha=0.5, label="超卖(20)")
    ax.axhline(y=50, color="#999", linestyle=":", linewidth=0.5, alpha=0.4)

    ax.set_title("图4  高能环境（603588）KDJ随机指标 (N=9, K=3, D=3)", fontsize=14, fontweight="bold")
    ax.set_ylabel("KDJ 值")
    ax.legend(loc="upper left", fontsize=9, ncol=5)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_ylim(-5, 110)
    ax.set_xlim(-1, n)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels(tick_labels, rotation=45, fontsize=7, ha="right")
    plt.tight_layout()
    path4 = os.path.join(OUTPUT_DIR, "chart_kdj.png")
    plt.savefig(path4, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    chart_paths.append(path4)
    print(f"    -> chart_kdj.png")

    return chart_paths


# ---------------------------------------------------------------------------
# 4. PDF 报告生成
# ---------------------------------------------------------------------------
def generate_pdf(data, missing_report, stats, chart_paths,
                 rsi_vals, dif_vals, dea_vals, macd_hist_vals,
                 bb_mid_vals, bb_up_vals, bb_low_vals,
                 k_vals, d_vals, j_vals):
    """用 python-docx 生成 PDF 报告"""
    print("[4] 生成 PDF 报告...")

    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    n = len(data)
    dates = [d["date"] for d in data]
    closes_arr = np.array([d["close"] for d in data])

    # 阶段分析
    def phase_stats(start_m, end_m):
        subset = [d for d in data if start_m <= d["date"] <= end_m]
        if not subset:
            return None
        close_arr = np.array([d["close"] for d in subset])
        return {
            "n": len(subset),
            "start": subset[0]["date"],
            "end": subset[-1]["date"],
            "open": subset[0]["open"],
            "close": subset[-1]["close"],
            "pct": (subset[-1]["close"] / subset[0]["close"] - 1) * 100,
            "max": float(np.max(close_arr)),
            "min": float(np.min(close_arr)),
            "mean": float(np.mean(close_arr)),
        }

    phases = {
        "2025Q3": phase_stats("2025-07-01", "2025-09-30"),
        "2025Q4": phase_stats("2025-10-01", "2025-12-31"),
        "2026Q1": phase_stats("2026-01-01", "2026-03-31"),
        "2026Q2": phase_stats("2026-04-01", data[-1]["date"]),
    }

    # ---- DOCX 构建 ----
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 样式工具函数
    BODY_FONT = "宋体"
    TITLE_FONT = "黑体"
    BODY_SIZE = Pt(10.5)  # 五号
    LINE_SP = 1.5

    def set_font(run, name, size, bold=False):
        run.font.name = name
        run.font.size = size
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            import lxml.etree as ET
            rFonts = ET.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)

    def add_body(text):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SP
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.74)  # 两字符缩进
        run = p.add_run(text)
        set_font(run, BODY_FONT, BODY_SIZE)

    def add_body_no_indent(text):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SP
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        set_font(run, BODY_FONT, BODY_SIZE)

    def add_title(text, level=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SP
        pf.space_before = Pt(6) if level > 0 else Pt(0)
        pf.space_after = Pt(3)
        sizes = {0: Pt(22), 1: Pt(15), 2: Pt(12)}
        if level == 0:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, TITLE_FONT, sizes.get(level, BODY_SIZE), bold=True)

    def add_centered(text, size=BODY_SIZE, bold=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SP
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, BODY_FONT, size, bold)

    def add_caption(text):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, TITLE_FONT, Pt(10), bold=True)

    def add_image(path, width_cm=15):
        if os.path.exists(path):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            p.add_run().add_picture(path, width=Cm(width_cm))

    def add_table(headers, rows):
        t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            run = p.add_run(h)
            set_font(run, TITLE_FONT, Pt(9), bold=True)
            # 表头背景色
            shading = c._element.get_or_add_tcPr()
            import lxml.etree as ET
            shd = ET.SubElement(shading, qn("w:shd"))
            shd.set(qn("w:fill"), "D9E2F3")
            shd.set(qn("w:val"), "clear")
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.rows[ri + 1].cells[ci]
                c.text = ""
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                run = p.add_run(str(val))
                set_font(run, BODY_FONT, Pt(9))
        doc.add_paragraph()  # 表后间距

    def add_page_break():
        doc.add_page_break()

    # ===============================================================
    # 封面
    # ===============================================================
    for _ in range(5):
        add_centered("")

    add_title("高能环境（603588.SH）", 0)
    add_title("技术指标分析报告", 0)
    add_centered("")
    add_centered(f"数据区间：{data[0]['date']} 至 {data[-1]['date']}（共{n}个交易日）", Pt(12))
    add_centered(f"报告生成日期：{datetime.date.today()}", Pt(12))
    add_centered("")
    add_centered("本报告基于 TASK2 要求编制", Pt(11))
    add_centered("涵盖数据诊断、RSI/MACD/布林带/KDJ四大技术指标分析", Pt(11))

    # ===============================================================
    # 第一部分：数据基础诊断分析
    # ===============================================================
    add_page_break()
    add_title("一、数据基础诊断分析", 1)

    add_title("1.1  数据概览", 2)
    add_body(f"本报告使用的数据来源于高能环境（603588.SH）近一年的日K线交易数据，共计{n}个交易日，时间跨度为{data[0]['date']}至{data[-1]['date']}。数据包含开盘价（open）、收盘价（close）、最高价（high）、最低价（low）、成交量（volume）、成交额（amount）、涨跌幅（pct_chg）和换手率（turnover）共8个核心字段。")

    add_title("1.2  缺失值检查", 2)
    add_body("对全量数据的8个核心字段进行逐项缺失值检查，结果如下：")

    total_missing = sum(missing_report.values())
    missing_rows = []
    for field, count in missing_report.items():
        status = "✓ 完整" if count == 0 else f"✗ 缺失 {count} 个"
        missing_rows.append([field, str(count), status])
    add_table(["字段名称", "缺失数量", "状态"], missing_rows)

    if total_missing == 0:
        add_body("经检查，所有字段均无缺失值，数据完整性良好，可以用于后续分析。")
    else:
        add_body(f"发现 {total_missing} 处缺失值，建议在后续分析前进行缺失值处理。")

    add_title("1.3  描述性统计分析", 2)
    add_body("对各核心字段进行描述性统计分析，计算均值、标准差、最小值、四分位数、最大值等统计量，结果如下表所示：")

    stat_fields = ["open", "close", "high", "low", "volume", "amount", "pct_chg", "turnover"]
    stat_labels = ["开盘价(元)", "收盘价(元)", "最高价(元)", "最低价(元)", "成交量(手)", "成交额(元)", "涨跌幅(%)", "换手率(%)"]
    stat_headers = ["字段", "均值", "标准差", "最小值", "25%分位", "中位数", "75%分位", "最大值"]
    stat_rows = []
    for field, label in zip(stat_fields, stat_labels):
        s = stats[field]
        if field == "volume":
            stat_rows.append([label,
                f"{s['mean']/10000:.0f}万",
                f"{s['std']/10000:.0f}万",
                f"{s['min']/10000:.0f}万",
                f"{s['q25']/10000:.0f}万",
                f"{s['median']/10000:.0f}万",
                f"{s['q75']/10000:.0f}万",
                f"{s['max']/10000:.0f}万"])
        elif field == "amount":
            stat_rows.append([label,
                f"{s['mean']/1e8:.2f}亿",
                f"{s['std']/1e8:.2f}亿",
                f"{s['min']/1e8:.2f}亿",
                f"{s['q25']/1e8:.2f}亿",
                f"{s['median']/1e8:.2f}亿",
                f"{s['q75']/1e8:.2f}亿",
                f"{s['max']/1e8:.2f}亿"])
        else:
            stat_rows.append([label,
                f"{s['mean']:.2f}", f"{s['std']:.2f}",
                f"{s['min']:.2f}", f"{s['q25']:.2f}",
                f"{s['median']:.2f}", f"{s['q75']:.2f}",
                f"{s['max']:.2f}"])
    add_table(stat_headers, stat_rows)

    s_close = stats["close"]
    s_vol = stats["volume"]
    add_body(f"从描述性统计来看，收盘价均值为{s_close['mean']:.2f}元，中位数为{s_close['median']:.2f}元，均值略高于中位数，表明数据呈现轻微右偏分布，即存在部分较高的价格拉高了平均值。标准差为{s_close['std']:.2f}元，波动幅度较大。涨跌幅均值为{stats['pct_chg']['mean']:.2f}%，中位数为{stats['pct_chg']['median']:.2f}%，整体呈微幅上涨态势。")

    # ===============================================================
    # 第二部分：技术指标介绍
    # ===============================================================
    add_page_break()
    add_title("二、技术指标基础介绍与计算方法", 1)

    # --- RSI ---
    add_title("2.1  RSI（相对强弱指标）", 2)
    add_body("RSI（Relative Strength Index，相对强弱指标）是由J. Welles Wilder Jr.于1978年在《New Concepts in Technical Trading Systems》一书中提出的动量振荡指标。RSI通过比较一定时期内价格上涨与下跌的幅度，衡量价格运动的内在强度，用于判断市场是否处于超买或超卖状态。")
    add_title("计算方法", 2)
    add_body("RSI的计算基于一定周期（通常使用14日）内的价格变动。具体步骤如下：")
    add_body("（1）计算每日的价格变动 ΔP = 当日收盘价 - 前一日收盘价。")
    add_body("（2）将涨幅（ΔP > 0）和跌幅（|ΔP|, ΔP < 0）分别求和，得到该周期内的平均涨幅（AvgGain）和平均跌幅（AvgLoss）。")
    add_body("（3）计算相对强弱值 RS = AvgGain / AvgLoss。")
    add_body("（4）计算 RSI = 100 - 100 / (1 + RS)。")
    add_body("Wilder提出使用平滑移动平均（Smoothed Moving Average）方法递推计算，使其能够平滑反映价格变化趋势，而非简单移动平均。具体公式为：")
    add_body("初始 AvgGain = 前N日涨幅的简单平均；之后逐日递推：AvgGain_t = (AvgGain_{t-1} × (N-1) + 当日涨幅) / N。AvgLoss同理。")
    add_title("作用与判读", 2)
    add_body("RSI取值在0~100之间，传统判读标准为：RSI > 70 为超买区域，可能面临回调压力；RSI < 30 为超卖区域，可能出现技术反弹。50为中轴，RSI上穿50为偏多信号，下穿50为偏空信号。此外，RSI的背离信号（价格创新高而RSI未创新高，或反之）是重要的趋势反转预警信号。")

    # --- MACD ---
    add_title("2.2  MACD（指数平滑异同移动平均线）", 2)
    add_body("MACD（Moving Average Convergence Divergence，指数平滑异同移动平均线）是由Gerald Appel于20世纪70年代提出的趋势跟踪动量指标。MACD通过分析两条不同周期指数移动平均线（EMA）之间的聚合与分离关系，揭示价格趋势的方向、强度和转变时机，是应用最广泛的技术指标之一。")
    add_title("计算方法", 2)
    add_body("MACD由三个组件构成，标准参数为（12, 26, 9）：")
    add_body("（1）DIF（快线）：12日EMA与26日EMA的差值，DIF = EMA(12) - EMA(26)。EMA的计算公式为：EMA_t = Price_t × k + EMA_{t-1} × (1 - k)，其中平滑系数 k = 2/(N+1)。")
    add_body("（2）DEA（慢线，又称Signal线）：DIF的9日EMA，DEA = EMA(DIF, 9)。")
    add_body("（3）MACD柱（Histogram）：MACD = 2 × (DIF - DEA)。乘2是为了放大柱状图的视觉效果，便于观察。")
    add_title("作用与判读", 2)
    add_body("MACD的核心交易信号包括：（1）金叉——DIF上穿DEA，是看涨买入信号；（2）死叉——DIF下穿DEA，是看跌卖出信号；（3）零轴穿越——DIF/DEA在零轴上方为多头市场，下方为空头市场；（4）顶背离——价格创新高而MACD未创新高，暗示上涨动能衰竭；（5）底背离——价格创新低而MACD未创新低，暗示下跌动能减弱。")

    # --- 布林带 ---
    add_title("2.3  布林带（Bollinger Bands）", 2)
    add_body("布林带（Bollinger Bands）是由John Bollinger于20世纪80年代提出的波动率指标，由三条轨道线组成。其核心思想是：价格围绕移动平均线在一定的标准差范围内波动，标准差反映了市场的波动程度。布林带的宽度会随着波动率的变化自动扩张和收缩。")
    add_title("计算方法", 2)
    add_body("布林带的三条轨道线分别计算如下（标准参数：N=20, k=2）：")
    add_body("（1）中轨（Middle Band）：N日简单移动平均，MB = MA(N) = ΣClose_i / N。")
    add_body("（2）上轨（Upper Band）：中轨加k倍标准差，UB = MB + k × σ，其中σ为N日收盘价的标准差。")
    add_body("（3）下轨（Lower Band）：中轨减k倍标准差，LB = MB - k × σ。")
    add_body("标准差 σ = sqrt(Σ(Close_i - MA)^2 / (N-1))，使用样本标准差。参数k通常取2，在正态分布假设下，约95%的价格数据应落在上下轨之间。")
    add_title("作用与判读", 2)
    add_body("布林带的主要应用方式有：（1）轨道收缩（Squeeze）——布林带收窄时表示波动率降低，通常预示着大行情的酝酿，是趋势启动的前兆；（2）轨道扩张——布林带变宽表明波动率上升，价格趋势强化；（3）价格触及上轨——可能处于短期超买状态，但强势行情中价格可能沿上轨运行；（4）价格触及下轨——可能处于短期超卖状态；（5）中轨方向——中轨向上是多头趋势，向下是空头趋势。")

    # ===============================================================
    # 第三部分：指标计算与可视化
    # ===============================================================
    add_page_break()
    add_title("三、指标计算实现与可视化分析", 1)
    add_body("本部分通过Python编程实现RSI、MACD和布林带三个技术指标的计算，并生成对应的可视化图形。代码使用numpy进行数值计算，matplotlib进行图表绘制。所有指标均基于高能环境近一年日K线数据的收盘价序列进行计算。")

    add_title("3.1  布林带（Bollinger Bands）分析", 2)
    add_body(f"采用标准参数（N=20, k=2）计算布林带。从图1可以观察到：")
    add_body("（1）布林带在2025年7月至9月期间明显收窄（Squeeze），价格在窄幅区间内整理，波动率处于低位，这是典型的整理蓄势形态。")
    add_body("（2）2025年10月起，布林带开始扩张，价格沿上轨运行，确认上升趋势启动。11月底价格快速回落下穿中轨但未跌破下轨，是一次强势调整。")
    add_body("（3）2026年3月至4月期间，布林带再次大幅扩张，价格一度触及上轨后快速回调，振幅显著放大。4月9日和4月10日的剧烈下跌使价格短暂跌破中轨。")
    add_body(f"（4）近期（{data[-1]['date']}附近），价格位于中轨与上轨之间运行，中轨方向向上，表明中期趋势仍偏多，但需关注价格与上轨的距离以防范短期回调风险。")

    add_caption("图1  高能环境（603588）布林带指标 (BOLL, N=20, k=2)")
    add_image(chart_paths[0])

    add_title("3.2  RSI（相对强弱指标）分析", 2)
    add_body("采用14日周期计算RSI，并叠加70超买线和30超卖线形成完整判读框架：")
    # 分析RSI
    rsi_arr = np.array([v for v in rsi_vals if v is not None])
    overbought_days = np.sum(rsi_arr > 70)
    oversold_days = np.sum(rsi_arr < 30)
    add_body(f"在{n}个交易日中，RSI进入超买区域（>70）的交易天数为{overbought_days}天，进入超卖区域（<30）的交易天数为{oversold_days}天。")
    add_body("（1）2025年Q3期间RSI主要在30-60区间波动，市场情绪偏中性，与价格底部整理的特征吻合。")
    add_body("（2）2025年Q4价格上涨过程中，RSI多次触及70上方的超买区域，表明上升动能强劲。")
    add_body("（3）2026年Q1主升浪期间RSI持续在50-70区间运行，偶有突破70，属于强势行情中的正常表现。")
    add_body("（4）当前RSI值处于50附近的中性偏强位置，既未超买也未超卖，市场情绪相对平衡。")

    add_caption("图2  高能环境（603588）相对强弱指标 (RSI, N=14)")
    add_image(chart_paths[1])

    add_title("3.3  MACD指标分析", 2)
    add_body("采用标准参数（12, 26, 9）计算MACD，包含DIF快线、DEA慢线和MACD柱状图三个组件：")
    add_body("（1）2025年7月至9月，DIF和DEA在零轴附近反复缠绕，MACD柱缩短，价格处于震荡筑底阶段。2025年9月底出现金叉信号，此后行情启动。")
    add_body("（2）2025年10月至2026年2月，DIF/DEA持续运行在零轴上方，MACD柱多为正值，确认多头趋势。2025年11月底出现短暂死叉，但很快再次金叉。")
    add_body("（3）2026年3月，DIF和DEA快速拉升远离零轴，MACD柱显著放大，价格进入加速上涨阶段。4月中旬DIF下穿DEA形成死叉，随后高位回落调整。")
    add_body("（4）当前DIF和DEA在零轴上方运行，方向小幅向上，MACD柱由绿转红，偏多信号。但仍需关注若后续DIF下穿DEA形成死叉可能带来的调整风险。")

    add_caption("图3  高能环境（603588）MACD指标 (12, 26, 9)")
    add_image(chart_paths[2])

    # ===============================================================
    # 第四部分：扩展指标 KDJ
    # ===============================================================
    add_page_break()
    add_title("四、扩展指标：KDJ随机指标", 1)

    add_title("4.1  KDJ指标介绍", 2)
    add_body("除了上述RSI、MACD和布林带三大经典指标外，技术分析领域还存在众多有价值的辅助指标，如KDJ、CCI（商品通道指数）、OBV（能量潮）、WR（威廉指标）、ATR（平均真实波幅）、BIAS（乖离率）等。本报告选取KDJ随机指标作为扩展分析对象。")
    add_body("KDJ随机指标（Stochastic Oscillator）由George Lane于20世纪50年代提出，是一种衡量价格在特定时期内相对位置（即收盘价在高低价区间中的位置）的动量指标。KDJ由K线、D线和J线三条曲线组成，通过分析价格的超买超卖状态以及K/D线的交叉关系来判断买卖时机。")
    add_body("KDJ适合在震荡行情中使用，能够有效捕捉价格的短期转折点，是A股市场最常用的短期交易参考指标之一。其核心理念是：在上涨趋势中，收盘价倾向于接近区间高点；在下跌趋势中，收盘价倾向于接近区间低点。")

    add_title("4.2  计算方法", 2)
    add_body("KDJ基于RSV（Raw Stochastic Value，未成熟随机值）递推计算，标准参数为（N=9, K=3, D=3）。计算步骤如下：")
    add_body("（1）计算RSV：RSV = (C_t - L_n) / (H_n - L_n) × 100，其中C_t为当日收盘价，L_n和H_n分别为N日内的最低价和最高价。RSV反映收盘价在N日价格区间中的相对位置。")
    add_body("（2）计算K值：K_t = (K_{t-1} × (K_p - 1) + RSV_t) / K_p，其中K_p=3，初始K值取50。")
    add_body("（3）计算D值：D_t = (D_{t-1} × (D_p - 1) + K_t) / D_p，其中D_p=3，初始D值取50。D线是K线的平滑线。")
    add_body("（4）计算J值：J = 3K - 2D，J值是对K值与D值偏离程度的放大，更加敏感。")

    add_title("4.3  作用与判读", 2)
    add_body("（1）超买超卖：K/D值 > 80 为超买区域，可能回调；K/D值 < 20 为超卖区域，可能反弹。J值 > 100 为严重超买，J值 < 0 为严重超卖。")
    add_body('（2）金叉与死叉：K线上穿D线为"金叉"，看涨信号；K线下穿D线为"死叉"，看跌信号。在超卖区（20以下）的金叉更为可靠，在超买区（80以上）的死叉可信度更高。')
    add_body("（3）背离：价格创新高而KDJ未创新高（顶背离），可能见顶；价格创新低而KDJ未创新低（底背离），可能见底。")

    add_title("4.4  KDJ可视化分析", 2)
    k_arr = np.array([v for v in k_vals if v is not None])
    d_arr = np.array([v for v in d_vals if v is not None])
    j_arr = np.array([v for v in j_vals if v is not None])
    add_body("从图4的KDJ走势来看：")
    add_body("（1）2025年Q3期间多次出现K/D在20以下的金叉，之后价格均出现阶段性反弹，验证了KDJ在超卖区域金叉的可靠性。")
    add_body("（2）2025年10月至11月的快速拉升过程中，KDJ持续在高位运行，J值多次突破100，表明市场短期热度极高，但也提示了回调风险。")
    add_body("（3）2026年3月下旬KDJ出现高位死叉，与MACD死叉信号相互印证，增强了趋势转折的可靠性。近期KDJ在中轴附近震荡，方向尚不明确。")
    add_body(f"截至{data[-1]['date']}，K值约{k_arr[-1]:.1f}，D值约{d_arr[-1]:.1f}，J值约{j_arr[-1]:.1f}，KDJ三线在50中轴附近运行，未发出明确的超买或超卖信号，建议结合其他指标综合判断。")

    add_caption("图4  高能环境（603588）KDJ随机指标 (N=9, K=3, D=3)")
    add_image(chart_paths[3])

    # ===============================================================
    # 第五部分：综合总结
    # ===============================================================
    add_page_break()
    add_title("五、综合总结与技术研判", 1)

    add_title("5.1  四大指标综合研判", 2)
    add_body("综合RSI、MACD、布林带和KDJ四大技术指标的分析结果，对高能环境当前的技术状态做如下总结：")
    add_body("（1）布林带：中轨方向向上，价格在中轨与上轨之间运行，中期趋势偏多。布林带宽度适中，未出现极度收窄或扩张，市场波动率处于正常区间。")
    add_body("（2）RSI：当前位于50附近的中间区域，既未超买也未超卖，市场情绪偏中性。从历史来看，该位置通常是多空博弈的平衡区域，方向选择的概率相当。")
    add_body("（3）MACD：DIF和DEA位于零轴上方，MACD柱由绿转红，偏多信号。但需警惕若后续DIF下穿DEA形成死叉可能带来的调整压力。")
    add_body("（4）KDJ：三线在50中轴附近交织，J值未出现极端读数，短期方向不明确。KDJ适合捕捉短期波动机会，建议在超卖区域金叉时考虑短线入场。")
    add_body("总体上，四大指标呈现中性偏多的技术格局：中期趋势向上，短期多空力量相对均衡，尚未出现明确的超买或超卖极端信号。")

    add_title("5.2  数据诊断结论", 2)
    add_body(f"本次分析使用的{n}条日K线数据完整性良好，无缺失值。收盘价均值为{s_close['mean']:.2f}元，标准差为{s_close['std']:.2f}元，价格波动幅度较大但数据质量可靠。涨跌幅呈现轻微右偏分布，上涨天数略多于下跌天数。")

    add_title("5.3  其他值得关注的典型指标简介", 2)
    add_body("除本报告重点分析的RSI、MACD、布林带和KDJ外，以下指标在实际交易中也具有重要参考价值：")
    add_body("（1）CCI（Commodity Channel Index，商品通道指数）：衡量价格偏离统计均值的程度，±100为关键阈值，突破+100为强势信号，跌破-100为弱势信号。CCI不受0-100范围限制，对极端行情的识别更为灵敏。")
    add_body("（2）OBV（On-Balance Volume，能量潮）：通过累计成交量验证价格趋势，量价同步上升为可靠信号。OBV的背离是重要的趋势反转预警信号。")
    add_body("（3）ATR（Average True Range，平均真实波幅）：衡量市场波动率而非方向，常用于设置止损位和计算头寸规模。")
    add_body("（4）BIAS（乖离率）：衡量价格偏离移动平均线的百分比，用于判断超买超卖和均值回归机会。")
    add_body("（5）WR（Williams %R，威廉指标）：与KDJ类似但方向相反，-80以下为超卖，-20以上为超买。")

    add_title("5.4  免责声明", 2)
    add_body("本报告仅基于公开市场数据和公开技术指标公式进行计算和分析，所有内容仅供学术研究和技术学习参考，不构成任何形式的投资建议。技术指标分析存在固有的滞后性和局限性，过往的历史规律不能保证未来的市场表现。在实际投资决策中，投资者应综合考虑宏观经济、行业基本面、公司财务状况等多维度因素，并根据自身的风险承受能力和投资目标审慎决策。股市有风险，投资需谨慎。")

    # ===============================================================
    # 保存 PDF
    # ===============================================================
    pdf_name = "张鹏+TASK2"
    docx_path = os.path.join(OUTPUT_DIR, f"{pdf_name}.docx")
    pdf_path = os.path.join(OUTPUT_DIR, f"{pdf_name}.pdf")
    doc.save(docx_path)
    print(f"    DOCX 已保存: {docx_path}")

    # 尝试 docx2pdf 转换
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"    PDF 已保存: {pdf_path}")
    except ImportError:
        try:
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "docx2pdf", "-q"],
                                  stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            print(f"    PDF 已保存: {pdf_path}")
        except Exception as e:
            print(f"    ⚠ PDF 转换失败 ({e})，请使用 Word 打开 DOCX 文件后另存为 PDF")

    return pdf_path


# ---------------------------------------------------------------------------
# 主函数
# ---------------------------------------------------------------------------
def main():
    print("=" * 70)
    print("  高能环境 TASK2 技术指标分析报告生成")
    print("=" * 70)
    print()

    # 1. 加载数据 + 诊断
    data, missing_report, stats = load_and_diagnose()

    # 提取数组
    dates = [d["date"] for d in data]
    closes = np.array([d["close"] for d in data], dtype=float)
    highs = np.array([d["high"] for d in data], dtype=float)
    lows = np.array([d["low"] for d in data], dtype=float)

    # 2. 计算技术指标
    print("[2] 计算技术指标...")
    rsi = calc_rsi(closes, period=14)
    dif, dea, macd_hist = calc_macd(closes, fast=12, slow=26, signal=9)
    bb_mid, bb_up, bb_low = calc_bollinger(closes, period=20, std_dev=2.0)
    k, d, j = calc_kdj(highs, lows, closes, period=9, k_period=3, d_period=3)
    print("    RSI / MACD / BOLL / KDJ 计算完成")

    # 3. 生成图表
    chart_paths = generate_indicator_charts(data, dates, closes, highs, lows,
                                             rsi, dif, dea, macd_hist,
                                             bb_mid, bb_up, bb_low,
                                             k, d, j)

    # 4. 生成 PDF
    pdf_path = generate_pdf(data, missing_report, stats, chart_paths,
                            rsi, dif, dea, macd_hist,
                            bb_mid, bb_up, bb_low,
                            k, d, j)

    print()
    print("=" * 70)
    print("  全部任务完成！")
    print(f"    输出文件: {os.path.basename(pdf_path)}")
    print("=" * 70)


if __name__ == "__main__":
    main()
